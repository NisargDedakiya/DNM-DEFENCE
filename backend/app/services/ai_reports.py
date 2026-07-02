"""
Module 5 — AI Report Generation Engine.

Feature 5.1: monthly report generator (exec summary via Claude + technical
             findings table + PDF/Word export)
Feature 5.2: alert notification drafting for critical findings
Feature 5.3: weekly threat digest generator

The AI is only ever used for *writing* — summarizing and phrasing findings
that already exist in the database. It never decides severity, invents
findings, or is trusted with anything the human reviewer can't verify
against the underlying Finding rows before a report goes out.
"""
import base64
import logging
import os
import secrets
from datetime import datetime, timedelta

import anthropic
from docx import Document
from docx.shared import RGBColor
from jinja2 import Environment, FileSystemLoader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.models import Client, Finding, FindingStatus, Report, ReportType, MetricSnapshot
from app.services.risk_score import compute_risk_score, risk_band
from app.services.compliance import get_compliance_summary

logger = logging.getLogger(__name__)

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates", "reports")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "generated_reports")
os.makedirs(OUTPUT_DIR, exist_ok=True)

_jinja_env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))


def _claude_client() -> anthropic.Anthropic:
    if not settings.ANTHROPIC_API_KEY:
        raise RuntimeError("ANTHROPIC_API_KEY is not set — cannot generate AI report content.")
    return anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)


def generate_risk_trend_chart(snapshots: list[MetricSnapshot]) -> bytes | None:
    """
    Feature 5.1 — risk score trend chart. Draws a simple line chart with
    Pillow (already a dependency via qrcode[pil], so no new heavy library
    like matplotlib is needed for one chart) instead of leaving the report
    with only a single point-in-time score. Returns PNG bytes, or None if
    there isn't enough history yet to plot a trend.
    """
    if len(snapshots) < 2:
        return None

    from PIL import Image, ImageDraw

    width, height, pad = 700, 260, 40
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    scores = [s.risk_score or 0 for s in snapshots]
    max_score = max(100, max(scores))
    plot_w, plot_h = width - 2 * pad, height - 2 * pad

    # axes
    draw.line([(pad, pad), (pad, height - pad)], fill="#cccccc", width=1)
    draw.line([(pad, height - pad), (width - pad, height - pad)], fill="#cccccc", width=1)

    points = []
    n = len(scores)
    for i, score in enumerate(scores):
        x = pad + (i / (n - 1)) * plot_w if n > 1 else pad
        y = (height - pad) - (score / max_score) * plot_h
        points.append((x, y))

    draw.line(points, fill="#16213e", width=3)
    for x, y in points:
        draw.ellipse([x - 3, y - 3, x + 3, y + 3], fill="#16213e")

    from io import BytesIO
    buf = BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def gather_monthly_data(db: Session, client: Client, period_start: datetime, period_end: datetime) -> dict:
    """Feature 5.1 — pulls all findings/alerts/scans from the current month."""
    findings = (
        db.query(Finding)
        .filter(Finding.client_id == client.id, Finding.created_at >= period_start, Finding.created_at <= period_end)
        .order_by(Finding.cvss_score.desc())
        .all()
    )
    open_findings = db.query(Finding).filter(
        Finding.client_id == client.id, Finding.status.notin_([FindingStatus.resolved, FindingStatus.verified])
    ).all()

    counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    for f in open_findings:
        if f.severity.value in counts:
            counts[f.severity.value] += 1

    return {
        "findings_this_month": findings,
        "open_findings": open_findings,
        "counts": counts,
        "risk_score": compute_risk_score(counts),
    }


def generate_ai_remediation(finding_title: str, description: str, evidence: dict, tech_context: str = "") -> str:
    """
    Feature: AI-generated remediation, replacing the static per-issue
    template strings used elsewhere (vuln_scan.py, cspm.py). Those
    templates stay as the fallback default -- this function is meant to
    be called opportunistically to produce something more specific to
    the actual finding, not to replace the fallback entirely (if the AI
    call fails, the template still applies).
    """
    client_ai = _claude_client()
    prompt = f"""You are a security engineer writing remediation guidance for a specific finding. Be concrete and actionable -- name specific commands, config settings, or AWS/GCP/Azure console steps where relevant, not generic advice like "apply best practices."

Finding: {finding_title}
Description: {description or 'N/A'}
Technical evidence: {str(evidence)[:500]}
{f"Stack context: {tech_context}" if tech_context else ""}

Write 2-4 sentences of remediation steps a mid-level engineer could follow directly. No preamble, just the steps."""

    try:
        response = client_ai.messages.create(
            model=settings.ANTHROPIC_MODEL, max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        return "".join(block.text for block in response.content if block.type == "text").strip()
    except Exception as e:
        logger.error(f"AI remediation generation failed, caller should fall back to static template: {e}")
        return ""


def generate_client_risk_analysis(client: Client, data: dict) -> str:
    """
    Feature: client-specific risk analysis narrative -- goes beyond the
    exec summary's plain-English tone to give a more analytical read on
    WHY the current posture looks the way it does and what pattern to
    watch, meant for a technical stakeholder (security-aware CTO,
    not a non-technical founder -- see generate_executive_summary for that).
    """
    client_ai = _claude_client()
    counts = data["counts"]
    top_findings = [(f.title, f.severity.value, f.cvss_score) for f in data["open_findings"][:8]]

    prompt = f"""Write a technical risk analysis (5-7 sentences) for {client.name}, a {client.industry or 'technology'} company.

Open findings: {counts['critical']} critical, {counts['high']} high, {counts['medium']} medium, {counts['low']} low
Top findings by severity: {top_findings}

Cover: (1) what pattern or theme connects the top findings, if any (e.g. "several findings trace back to unpatched dependencies" or "exposure is concentrated in cloud misconfigurations rather than application code"), (2) which single finding poses the most realistic business risk and why, (3) a forward-looking note on what to prioritize next cycle. Write for a technical stakeholder who wants substance, not reassurance. No preamble."""

    response = client_ai.messages.create(
        model=settings.ANTHROPIC_MODEL, max_tokens=450,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(block.text for block in response.content if block.type == "text").strip()


def generate_executive_summary(client: Client, data: dict) -> str:
    """
    Feature 5.1 — plain-English exec summary for a non-technical founder/CTO.
    The prompt gives Claude only aggregate counts and top finding titles —
    never raw evidence/exploit detail — since this text may be shared outside
    the security team.
    """
    client_ai = _claude_client()
    counts = data["counts"]
    top_titles = [f.title for f in data["open_findings"][:5]]

    prompt = f"""You are writing the executive summary for a monthly security report for a non-technical startup founder/CTO. Do not use jargon.

Client: {client.name} ({client.industry or "technology company"})
Currently open findings: {counts['critical']} critical, {counts['high']} high, {counts['medium']} medium, {counts['low']} low
Top open issues: {', '.join(top_titles) if top_titles else 'None currently open'}
New findings discovered this reporting period: {len(data['findings_this_month'])}

Write a 3-4 sentence executive summary covering: (1) overall security posture in plain English, (2) the single most important thing they should know or do this month, (3) a brief note of encouragement/context if the trend is positive, or urgency if it's not. Do not include a greeting or sign-off — just the summary paragraph itself."""

    response = client_ai.messages.create(
        model=settings.ANTHROPIC_MODEL,
        max_tokens=400,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(block.text for block in response.content if block.type == "text").strip()


def generate_compliance_summary(db: Session, client: Client) -> str:
    """Feature 5.1 — compliance status paragraph, pulled from the real per-framework rollup (Module 6.4)."""
    summary = get_compliance_summary(db, client.id)
    if not summary:
        return (f"Compliance control tracking is not yet configured for {client.name}. "
                f"Set up SOC 2 / ISO 27001 / India DPDP Act checklists in the Compliance Center to populate this section.")

    framework_labels = {"soc2": "SOC 2", "iso27001": "ISO 27001", "india_dpdp": "India DPDP Act"}
    lines = []
    for fw, s in summary.items():
        label = framework_labels.get(fw, fw)
        lines.append(f"{label}: {s['percent_implemented']}% implemented ({s['implemented']}/{s['total']} controls; "
                     f"{s['in_progress']} in progress, {s['missing']} missing).")
    return " ".join(lines)


def render_report_html(db: Session, client: Client, data: dict, executive_summary: str, period_label: str,
                        risk_analysis: str = "", trend_chart_png: bytes | None = None) -> str:
    template = _jinja_env.get_template("monthly_report.html")
    chart_b64 = base64.b64encode(trend_chart_png).decode() if trend_chart_png else None
    return template.render(
        client_name=client.name,
        client_logo_url=client.logo_url,
        client_brand_color=client.brand_color or "#16213e",
        period_label=period_label,
        risk_score=data["risk_score"],
        risk_band=risk_band(data["risk_score"]),
        executive_summary=executive_summary,
        risk_analysis=risk_analysis,
        counts=data["counts"],
        findings=[{
            "severity": f.severity.value, "title": f.title, "cvss_score": f.cvss_score,
            "status": f.status.value, "created_at": f.created_at.strftime("%Y-%m-%d"),
        } for f in data["open_findings"]],
        compliance_summary=generate_compliance_summary(db, client),
        risk_trend_chart_b64=chart_b64,
        generated_at=datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
    )


def export_pdf(html_content: str, output_path: str) -> str:
    # Lazy import: weasyprint needs system libs (pango/cairo) not present in
    # every environment (e.g. CI), and nothing else in this module needs it.
    from weasyprint import HTML
    HTML(string=html_content).write_pdf(output_path)
    return output_path


def export_docx(db: Session, client: Client, data: dict, executive_summary: str, period_label: str, output_path: str,
                 risk_analysis: str = "", trend_chart_png: bytes | None = None) -> str:
    """Feature 5.1 — Word export for clients who need an editable report."""
    doc = Document()

    title = doc.add_heading(client.name, level=0)
    if client.brand_color:
        try:
            rgb = client.brand_color.lstrip("#")
            title.runs[0].font.color.rgb = RGBColor(int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16))
        except (ValueError, IndexError):
            pass  # malformed hex color -- fall back to default styling rather than fail the whole report
    doc.add_paragraph(f"Monthly Security Report — {period_label}").italic = True

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(executive_summary)

    if risk_analysis:
        doc.add_heading("Technical Risk Analysis", level=1)
        doc.add_paragraph(risk_analysis)

    doc.add_heading("This Month at a Glance", level=1)
    counts = data["counts"]
    doc.add_paragraph(f"Critical: {counts['critical']}  |  High: {counts['high']}  |  "
                       f"Medium: {counts['medium']}  |  Low: {counts['low']}")

    if trend_chart_png:
        doc.add_heading("Risk Score Trend", level=1)
        from io import BytesIO
        from docx.shared import Inches
        doc.add_picture(BytesIO(trend_chart_png), width=Inches(6))

    doc.add_heading("Technical Findings", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Severity", "Finding", "CVSS", "Status"
    for f in data["open_findings"]:
        row = table.add_row().cells
        row[0].text = f.severity.value.upper()
        row[1].text = f.title
        row[2].text = str(f.cvss_score)
        row[3].text = f.status.value

    doc.add_heading("Compliance Status", level=1)
    doc.add_paragraph(generate_compliance_summary(db, client))

    doc.save(output_path)
    return output_path


def generate_monthly_report(db: Session, client: Client) -> Report:
    """Orchestrates the full pipeline: gather data -> AI summary -> render -> export PDF+DOCX -> save Report row."""
    now = datetime.utcnow()
    period_start = (now.replace(day=1) - timedelta(days=1)).replace(day=1)
    period_end = now.replace(day=1) - timedelta(seconds=1)
    period_label = period_start.strftime("%B %Y")

    data = gather_monthly_data(db, client, period_start, period_end)
    exec_summary = generate_executive_summary(client, data)
    risk_analysis = generate_client_risk_analysis(client, data)

    # Feature 5.1 risk score trend chart — last 6 months of daily snapshots (Phase 1's MetricSnapshot table)
    trend_cutoff = now - timedelta(days=180)
    snapshots = (
        db.query(MetricSnapshot)
        .filter(MetricSnapshot.client_id == client.id, MetricSnapshot.snapshot_date >= trend_cutoff)
        .order_by(MetricSnapshot.snapshot_date.asc())
        .all()
    )
    trend_chart_png = generate_risk_trend_chart(snapshots)

    html = render_report_html(db, client, data, exec_summary, period_label, risk_analysis, trend_chart_png)

    safe_name = client.name.replace(" ", "_").replace("/", "-")
    pdf_path = os.path.join(OUTPUT_DIR, f"{safe_name}_{period_start.strftime('%Y-%m')}.pdf")
    docx_path = os.path.join(OUTPUT_DIR, f"{safe_name}_{period_start.strftime('%Y-%m')}.docx")
    export_pdf(html, pdf_path)
    export_docx(db, client, data, exec_summary, period_label, docx_path, risk_analysis, trend_chart_png)

    report = Report(
        client_id=client.id, report_type=ReportType.monthly_security,
        period_start=period_start, period_end=period_end,
        executive_summary=exec_summary, risk_analysis=risk_analysis, risk_score=data["risk_score"],
        pdf_path=pdf_path, docx_path=docx_path,
        share_token=secrets.token_urlsafe(24),  # Feature 6.5 — read-only share link
        created_at=now,
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    return report


def draft_alert_notification(finding: Finding) -> str:
    """
    Feature 5.2 — drafts a client-facing alert for a critical/high finding.
    Human review before send is the default (see Celery task); auto-send
    is opt-in per client, not the default here.
    """
    client_ai = _claude_client()
    prompt = f"""Draft a short client alert notification (email-length, 4-6 sentences) about this security finding:

Title: {finding.title}
Severity: {finding.severity.value}
Description: {finding.description or 'N/A'}
Remediation: {finding.remediation_steps or 'N/A'}

Structure: (1) what was found, in plain English, (2) why it matters / potential impact, (3) what the client should do immediately, (4) what the security team will do next. No greeting or sign-off, just the body."""

    response = client_ai.messages.create(
        model=settings.ANTHROPIC_MODEL, max_tokens=350,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(block.text for block in response.content if block.type == "text").strip()


def generate_weekly_threat_digest(client: Client, recent_finding_titles: list[str],
                                   cve_hits: list[dict] | None = None, threat_intel_hits: list[dict] | None = None) -> str:
    """
    Feature 5.3 — 1-page plain-English weekly digest, grounded in this
    week's actual scan data (CVE matches from Module 1.3's CIRCL lookup,
    threat-intel correlation from Module 3) rather than asking Claude to
    recall "current" CVEs/threats from its training data, which risks
    hallucinating dates, CVE IDs, and campaigns that aren't real or
    aren't current. If there's no real data this week, the digest says
    so honestly instead of inventing content to fill the space.
    """
    client_ai = _claude_client()
    cve_hits = cve_hits or []
    threat_intel_hits = threat_intel_hits or []

    cve_lines = [f"- {h.get('cve_id', 'CVE')}: {h.get('technology')} {h.get('version')} on {h.get('host')}" for h in cve_hits[:10]]
    intel_lines = [f"- {h.get('note', h)}" for h in threat_intel_hits[:10]]

    if not recent_finding_titles and not cve_lines and not intel_lines:
        source_material = "No new findings, CVE matches, or threat-intel hits this week — genuinely quiet."
    else:
        source_material = (
            f"New findings this week: {', '.join(recent_finding_titles) if recent_finding_titles else 'none'}\n"
            f"CVE matches found on their own infrastructure this week:\n{chr(10).join(cve_lines) if cve_lines else '  none'}\n"
            f"Threat-intel correlation hits this week:\n{chr(10).join(intel_lines) if intel_lines else '  none'}"
        )

    prompt = f"""Write a short (200-300 word) weekly security digest for a startup in the {client.industry or 'technology'} industry.

This is grounded strictly in real data collected from THIS CLIENT'S OWN infrastructure this week — do not invent, recall, or reference any CVEs, threats, or campaigns that aren't in the data below, even if you know of real current events. If the data below is thin, write a shorter, honest digest rather than padding it with general knowledge.

{source_material}

Cover: what was actually found this week (in plain English), why it matters, and one actionable takeaway. No jargon, no greeting/sign-off — just the digest body."""

    response = client_ai.messages.create(
        model=settings.ANTHROPIC_MODEL, max_tokens=500,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(block.text for block in response.content if block.type == "text").strip()
